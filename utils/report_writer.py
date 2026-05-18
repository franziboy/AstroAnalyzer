"""
AstroAnalyzer – Report-Writer
Erzeugt Markdown (.md), Word (.docx) und CSV-Ausgaben
"""

import csv, os, re
from datetime import datetime
from pathlib import Path


def _now():
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _date():
    return datetime.now().strftime("%Y-%m-%d")


# ── Markdown ───────────────────────────────────────────────────────────────────

def write_markdown(all_results: dict, output_path: str, meta: dict) -> str:
    """Schreibt vollständigen Markdown-Report."""
    r1 = all_results.get("m1", {})
    r2 = all_results.get("m2", {})
    r3 = all_results.get("m3", [])
    r4 = all_results.get("m4", {})
    r5 = all_results.get("m5", {})
    b1 = all_results.get("m1_bewertung", [])

    target  = meta.get("OBJECT", "Unbekannt")
    instr   = meta.get("INSTRUME", "?")
    tele    = meta.get("TELESCOP", "?")
    filt    = meta.get("FILTER", "?")
    exptime = meta.get("EXPTIME", 0)
    n_subs  = all_results.get("n_subs", meta.get("NCOMBINE", 1))
    total_h = round(n_subs * exptime / 3600, 2)

    lines = [
        f"---",
        f"datum: {_date()}",
        f"target: {target}",
        f"filter: {filt}",
        f"gesamtbelichtung: {total_h}h",
        f"setup: {tele} + {instr}",
        f"erstellt: AstroAnalyzer",
        f"tags: [astrofotografie, ergebnis]",
        f"---",
        f"",
        f"# Analyse: {target} – {_date()}",
        f"",
        f"## Setup",
        f"| Parameter | Wert |",
        f"|-----------|------|",
        f"| Teleskop / Brennweite | {tele} / {meta.get('FOCALLEN','?')} mm |",
        f"| Kamera | {instr} |",
        f"| Filter | {filt} |",
        f"| Belichtung | {n_subs} × {exptime:.0f} s = {total_h} h |",
        f"| Gain / Offset | {meta.get('GAIN','?')} / {meta.get('OFFSET','?')} |",
        f"| Sensortemperatur | {meta.get('CCD-TEMP','?')} °C |",
        f"| Bildskala | {r1.get('pixel_scale','?')} \"/px |",
        f"| Bildgröße | {meta.get('WIDTH','?')} × {meta.get('HEIGHT','?')} px |",
        f"",
        f"---",
        f"",
        f"## 1. Technische Bildqualität",
        f"",
        f"| Messgröße | Wert | Bewertung | Hinweis |",
        f"|-----------|------|-----------|---------|",
    ]

    for b in b1:
        lines.append(f"| {b['thema']} | {b['wert']} | {b['status']} | {b['hinweis']} |")

    _sky_str   = f"{r1['sky_median']:.6g}"     if r1.get('sky_median')    is not None else "?"
    _noise_str = f"{r1['noise_mad']:.3g}"      if r1.get('noise_mad')     is not None else "?"
    _grad_str  = f"{r1['gradient_pct']:.1f}"   if r1.get('gradient_pct')  is not None else "?"
    _gsig_str  = f"{r1['gradient_sigma']:.1f}" if r1.get('gradient_sigma') is not None else "?"
    _band_str  = f"{r1['banding_amp']:.3g}"    if r1.get('banding_amp')   is not None else "?"

    lines += [
        f"",
        f"### Detailmesswerte",
        f"",
        f"| Parameter | Wert |",
        f"|-----------|------|",
        f"| FWHM (median) | {r1.get('fwhm_median_px','?')} px = {r1.get('fwhm_arcsec','?')}\" |",
        f"| FWHM Zentrum / Rand | {r1.get('fwhm_center_px','?')} / {r1.get('fwhm_edge_px','?')} px |",
        f"| Feldvariation FWHM | Δ {r1.get('fwhm_field_delta','?')} px |",
        f"| Exzentrizität (med / p90) | {r1.get('ecc_median','?')} / {r1.get('ecc_p90','?')} |",
        f"| Hintergrund-Median | {_sky_str} |",
        f"| Rauschen (MAD) | {_noise_str} |",
        f"| Gradient | {_grad_str}% ({_gsig_str}σ) |",
        f"| Banding-Amplitude | {_band_str} |",
        f"| Gesättigte Pixel | {r1.get('sat_pixels','?')} |",
        f"| Vermessene Sterne | {r1.get('n_stars_measured','?')} |",
        f"",
        f"---",
        f"",
        f"## 2. Aufnahmetechnik",
        f"",
    ]

    # Integration
    intg = r2.get("integration", {})
    lines += [
        f"| Parameter | Wert | Bewertung |",
        f"|-----------|------|-----------|",
        f"| Gesamtbelichtung | {intg.get('total_h','?')} h | {intg.get('status','?')} |",
        f"| SNR-Gewinn | √{intg.get('n_subs','?')} = {intg.get('snr_gain','?')}× | |",
        f"| Tiefengewinn vs 1h | +{intg.get('dmag_vs_1h','?')} mag | |",
        f"| Sub-Länge | {r2.get('sub_laenge',{}).get('t_sub_s','?')} s | {r2.get('sub_laenge',{}).get('status','?')} |",
        f"| Gain-Regime | {r2.get('gain',{}).get('regime','?')} | {r2.get('gain',{}).get('status','?')} |",
        f"| Dithering | Pattern {r2.get('dithering',{}).get('ratio','?')}× Stack-Rauschen | {r2.get('dithering',{}).get('status','?')} |",
        f"| Temperatur | {r2.get('temperatur',{}).get('temp_c','?')} °C | {r2.get('temperatur',{}).get('status','?')} |",
        f"",
    ]
    rec2 = all_results.get("m2_empfehlungen", [])
    if rec2:
        lines.append("**Empfehlungen:**\n")
        for rec in rec2:
            lines.append(f"- {rec}")
        lines.append("")

    lines += [
        f"---",
        f"",
        f"## 3. PixInsight-Optimierung",
        f"",
    ]
    for step in r3:
        lines += [
            f"### {step['nr']}. {step['werkzeug']} _{step['phase']}_",
            f"",
            f"**Einstellung:** {step['einstellung']}",
            f"",
            f"> **Warum:** {step['begruendung']}",
            f"",
        ]

    lines += [
        f"---",
        f"",
        f"## 4. Kalibrierung",
        f"",
    ]
    dk = r4.get("dark", {})
    fl = r4.get("flat", {})
    fm = r4.get("flat_mismatch", {})

    _ag = f"{dk['amp_glow']:.2e}" if dk.get('amp_glow') else '—'
    _hp = dk.get('hot_px_pct')
    _hp_str = f"{_hp:.2f}%" if _hp is not None else "—"
    lines += [
        f"| Prüfung | Wert | Bewertung |",
        f"|---------|------|-----------|",
        f"| Dark: Amp-Glow | {_ag} | {dk.get('bewertung','—')} |",
        f"| Dark: Warmpixel | {_hp_str} | {'✅' if not dk.get('hot_px_warn') else '⚠️'} |",
        f"| Flat: Asymmetrie | {fl.get('asymmetrie_pct','—')}% | {fl.get('bewertung','—')} |",
        f"| Flat: Staubschatten | {fl.get('dust_amp_pct','—')}% | {'✅' if (fl.get('dust_amp_pct') or 0) < 1 else '⚠️'} |",
        f"| Flat-Mismatch corr | {fm.get('korrelation','—')} | {'⚠️ Mismatch!' if fm.get('warnung') else '✅ OK'} |",
        f"",
    ]
    if fm.get("interpretation"):
        lines += [f"**Befund:** {fm['interpretation']}", f""]

    lines += [
        f"---",
        f"",
        f"## 5. Hauptobjekte und Galaxienfeld",
        f"",
    ]
    if r5.get("wcs_used"):
        lines += [
            f"**Insgesamt im Feld:** {r5.get('n_ngc_ic',0)} NGC/IC · "
            f"{r5.get('n_high',0)} PGC-HIGH · "
            f"{r5.get('n_mittel',0)} PGC-MITTEL",
            f"",
            f"### Katalogisierte NGC/IC-Galaxien",
            f"",
            f"| Name | Typ | B-mag | Größe | RA | Dec |",
            f"|------|-----|-------|-------|----|-----|",
        ]
        for g in sorted(r5.get("ngc_ic", []), key=lambda x: x.get("B") or 99):
            lines.append(
                f"| {g['name']} | {g.get('hubble','-')} | "
                f"{'%.1f'%g['B'] if g.get('B') else '-'} | "
                f"{'%.2f'%g['maj']+'′' if g.get('maj') else '-'} | "
                f"{_ra_str(g.get('ra',0))} | {_dec_str(g.get('dec',0))} |"
            )
        lines.append("")
    else:
        lines.append("*Keine WCS-Lösung vorhanden — Galaxien-Koordinaten nicht verfügbar.*\n")

    # Objektbeschreibungen
    desc = r5.get("beschreibungen", {})
    if desc:
        lines += ["### Ausführliche Objektbeschreibungen", ""]
        for name, txt in desc.items():
            lines += [txt, ""]

    lines += [
        f"---",
        f"",
        f"*Erstellt mit AstroAnalyzer · {_now()}*",
    ]

    md = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md)
    return md


# ── Word (.docx) ───────────────────────────────────────────────────────────────

def write_docx(all_results: dict, output_path: str, meta: dict) -> None:
    """Schreibt vollständigen Word-Bericht."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    BLUE = RGBColor(0x1F, 0x4E, 0x79)

    # Stil-Helfer
    def h1(txt):
        p = doc.add_heading(txt, level=1)
        p.runs[0].font.color.rgb = BLUE
        return p

    def h2(txt):
        p = doc.add_heading(txt, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return p

    def body(txt):
        doc.add_paragraph(txt)

    def table_2col(rows, header=None):
        n_cols = len(header) if header else (len(rows[0]) if rows else 2)
        n_rows = len(rows) + (1 if header else 0)
        t = doc.add_table(rows=n_rows, cols=n_cols)
        t.style = "Table Grid"
        ri = 0
        if header:
            for ci, h in enumerate(header):
                cell = t.cell(0, ci)
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
            ri = 1
        for r_data in rows:
            for ci, val in enumerate(r_data):
                t.cell(ri, ci).text = str(val)
            ri += 1
        return t

    r1 = all_results.get("m1", {})
    r2 = all_results.get("m2", {})
    r3 = all_results.get("m3", [])
    r4 = all_results.get("m4", {})
    r5 = all_results.get("m5", {})
    b1 = all_results.get("m1_bewertung", [])
    target = meta.get("OBJECT", "Unbekannt")
    n_subs = all_results.get("n_subs", meta.get("NCOMBINE", 1))
    exptime = meta.get("EXPTIME", 0)
    total_h = round(n_subs * exptime / 3600, 2)

    # Titel
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(f"AstroAnalyzer – {target}")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = BLUE
    doc.add_paragraph(
        f"{meta.get('TELESCOP','?')} + {meta.get('INSTRUME','?')} · "
        f"{meta.get('FILTER','?')} · {total_h}h · {_date()}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Bildqualität
    h1("1. Technische Bildqualität")
    table_2col(
        [(b["thema"], f"{b['wert']}  {b['status']}  {b['hinweis']}") for b in b1],
        header=["Messgröße", "Ergebnis"]
    )

    # 2. Aufnahmetechnik
    h1("2. Aufnahmetechnik")
    intg = r2.get("integration", {})
    table_2col([
        ["Gesamtbelichtung",  f"{intg.get('total_h','?')} h   {intg.get('status','')}"],
        ["SNR-Gewinn",        f"√{intg.get('n_subs','?')} = {intg.get('snr_gain','?')}×"],
        ["Sub-Länge",         f"{r2.get('sub_laenge',{}).get('t_sub_s','?')} s   {r2.get('sub_laenge',{}).get('status','')}"],
        ["Gain/Regime",       f"{r2.get('gain',{}).get('gain','?')} — {r2.get('gain',{}).get('regime','?')}"],
        ["Dithering",         f"Pattern {r2.get('dithering',{}).get('ratio','?')}×   {r2.get('dithering',{}).get('status','')}"],
        ["Temperatur",        f"{r2.get('temperatur',{}).get('temp_c','?')} °C   {r2.get('temperatur',{}).get('status','')}"],
    ], header=["Parameter", "Wert / Bewertung"])
    for rec in all_results.get("m2_empfehlungen", []):
        doc.add_paragraph(rec, style="List Bullet")

    # 3. PixInsight
    h1("3. PixInsight-Optimierung")
    for step in r3:
        h2(f"{step['nr']}. {step['werkzeug']} [{step['phase']}]")
        body(f"Einstellung: {step['einstellung']}")
        p = doc.add_paragraph()
        p.add_run("Warum: ").bold = True
        p.add_run(step["begruendung"]).italic = True

    # 4. Kalibrierung
    h1("4. Kalibrierung")
    dk = r4.get("dark", {}); fl = r4.get("flat", {}); fm = r4.get("flat_mismatch", {})
    _ag = f"{dk['amp_glow']:.2e}" if dk.get('amp_glow') else '—'
    _hp = dk.get('hot_px_pct')
    _hp_str = f"{_hp:.2f}%" if _hp is not None else "—"
    table_2col([
        ["Dark: Amp-Glow",   f"{_ag}   {dk.get('bewertung','—')}"],
        ["Dark: Warmpixel",  _hp_str],
        ["Flat: Asymmetrie", f"{fl.get('asymmetrie_pct','—')}%   {fl.get('bewertung','—')}"],
        ["Flat-Mismatch",    f"corr = {fm.get('korrelation','—')}   {'⚠️' if fm.get('warnung') else '✅'}"],
    ], header=["Prüfung", "Ergebnis"])
    if fm.get("interpretation"):
        body(fm["interpretation"])

    # 5. Hauptobjekte
    h1("5. Hauptobjekte und Galaxienfeld")
    if r5.get("wcs_used"):
        body(f"NGC/IC: {r5.get('n_ngc_ic',0)} · PGC-HIGH: {r5.get('n_high',0)} · PGC-MITTEL: {r5.get('n_mittel',0)}")
        rows = [
            [g["name"], g.get("hubble","-"),
             f"{g['B']:.1f}" if g.get("B") else "-",
             f"{g['maj']:.2f}′" if g.get("maj") else "-",
             _ra_str(g.get("ra",0)), _dec_str(g.get("dec",0))]
            for g in sorted(r5.get("ngc_ic",[]), key=lambda x: x.get("B") or 99)
        ]
        table_2col(rows, header=["Name","Typ","B-mag","Größe","RA","Dec"])
        for name, txt in r5.get("beschreibungen", {}).items():
            h2(name)
            # Strip markdown
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", txt)
            body(clean)
    else:
        body("Keine WCS-Lösung — Galaxien-Koordinaten nicht verfügbar.")

    doc.save(output_path)


# ── CSV ────────────────────────────────────────────────────────────────────────

def write_csv(results_m5: dict, output_path: str) -> None:
    """Schreibt Galaxien-CSV."""
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Kategorie","Name","Typ","RA (J2000)","Dec (J2000)",
                    "B-mag","Größe (arcmin)","Pixel X","Pixel Y","Anmerkung"])
        for g in sorted(results_m5.get("ngc_ic",[]), key=lambda x: x.get("B") or 99):
            w.writerow([
                "NGC/IC", g["name"], g.get("hubble","-"),
                _ra_str(g.get("ra",0)), _dec_str(g.get("dec",0)),
                f"{g['B']:.1f}" if g.get("B") else "-",
                f"{g['maj']:.2f}" if g.get("maj") else "-",
                int(g["x"]), int(g["y"]), ""
            ])
        for p in results_m5.get("pgc_kand",[]):
            if p["konfidenz"] in ("HIGH","MITTEL"):
                w.writerow([
                    f"PGC-{p['konfidenz']}", "anonym", "?",
                    _ra_str(p.get("ra") or 0), _dec_str(p.get("dec") or 0),
                    "-","-", int(p["x"]), int(p["y"]), p["konfidenz"]
                ])


# ── Hilfsfunktionen ────────────────────────────────────────────────────────────

def _ra_str(ra_deg: float) -> str:
    h = ra_deg / 15
    hh = int(h); mm = int((h-hh)*60); ss = ((h-hh)*60-mm)*60
    return f"{hh:02d}h {mm:02d}m {ss:04.1f}s"

def _dec_str(dec_deg: float) -> str:
    sign = "+" if dec_deg >= 0 else "-"
    d = abs(dec_deg); dg = int(d); dm = int((d-dg)*60); ds = ((d-dg)*60-dm)*60
    return f"{sign}{dg:02d}° {dm:02d}′ {ds:04.1f}″"
